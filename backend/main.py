from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
import whisper
from docx import Document
import os
import asyncio
from typing import Dict
import uuid
from datetime import datetime, timedelta
from pydub import AudioSegment
from moviepy.editor import AudioFileClip
import torch
import glob

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Store job status
JOBS: Dict[str, Dict] = {}


def update_job_status(
    job_id: str, stage: str, current: int = 0, total_chunks: int = 0, message: str = ""
):
    """Update job status with detailed progress information"""
    if job_id in JOBS:
        JOBS[job_id].update(
            {
                "status": "processing",
                "progress": {
                    "stage": stage,
                    "current_chunk": current,
                    "total_chunks": total_chunks,
                    "message": message,
                },
            }
        )


def cleanup_old_jobs():
    """Remove jobs older than 24 hours"""
    now = datetime.now()
    for job_id in list(JOBS.keys()):
        if now - JOBS[job_id]["timestamp"] > timedelta(hours=24):
            del JOBS[job_id]


async def split_audio(file_path: str, chunk_duration: int = 120000) -> list:
    """Split audio file into chunks"""
    audio = AudioSegment.from_file(file_path)
    chunks = []

    # Optimize audio before splitting
    audio = audio.set_channels(1)  # Convert to mono
    audio = audio.set_frame_rate(16000)  # Reduce sample rate

    for i in range(0, len(audio), chunk_duration):
        chunk = audio[i : i + chunk_duration]
        chunk_path = f"{file_path}_chunk_{i//chunk_duration}.wav"
        chunk.export(
            chunk_path,
            format="wav",
            parameters=["-q:a", "0", "-ar", "16000"],  # Optimize export
        )
        chunks.append(chunk_path)
        del chunk  # Free memory

    del audio  # Free memory
    return chunks


async def convert_to_wav(input_path: str) -> str:
    """Convert audio file to WAV format"""
    try:
        output_path = f"{os.path.splitext(input_path)[0]}.wav"

        # Use pydub for conversion
        audio = AudioSegment.from_file(input_path)
        audio = audio.set_channels(1)  # Convert to mono
        audio = audio.set_frame_rate(16000)  # Set sample rate

        # Export as WAV with optimal parameters
        audio.export(
            output_path, format="wav", parameters=["-q:a", "0", "-ar", "16000"]
        )

        # Clean up original file
        os.remove(input_path)
        return output_path

    except Exception as e:
        if os.path.exists(output_path):
            os.remove(output_path)
        raise Exception(f"Errore nella conversione audio: {str(e)}")


async def process_audio(job_id: str, file_path: str, output_path: str):
    try:
        # Initialize variables
        current_step = 1
        total_steps = 6
        transcriptions = []

        # Ensure absolute paths
        file_path = os.path.abspath(file_path)
        output_path = os.path.abspath(output_path)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        temp_dir = os.path.join(base_dir, "temp")
        output_dir = os.path.join(base_dir, "output")

        # Create directories if they don't exist
        os.makedirs(temp_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File audio non trovato: {file_path}")

        # Initial validation
        update_job_status(
            job_id,
            "validating",
            current=current_step,
            total_chunks=total_steps,
            message="Fase 1/6: Validazione del file audio..."
        )

        current_step += 1

        # GPU Check
        current_step += 1  # Increment before GPU check
        if torch.cuda.is_available():
            gpu_name = torch.cuda.get_device_name(0)
            update_job_status(
                job_id,
                "gpu_check",
                current=current_step,
                total_chunks=total_steps,
                message=f"Fase 2/6: GPU rilevata ({gpu_name}). Ottimizzazione in corso..."
            )
        else:
            update_job_status(
                job_id,
                "gpu_check",
                current=current_step,
                total_chunks=total_steps,
                message="Fase 2/6: GPU non disponibile. Utilizzo CPU (processo più lento)..."
            )

        # Audio conversion
        update_job_status(
            job_id,
            "converting",
            current=current_step,
            total_chunks=total_steps,
            message="Fase 3/6: Ottimizzazione del file audio..."
        )
        wav_path = await convert_to_wav(file_path)

        # Model loading
        update_job_status(
            job_id,
            "loading_model",
            current=current_step,
            total_chunks=total_steps,
            message="Fase 4/6: Caricamento del modello di intelligenza artificiale..."
        )
        model = await load_model()  # New function for model loading

        # Audio splitting
        update_job_status(
            job_id,
            "splitting_audio",
            current=current_step,
            total_chunks=total_steps,
            message="Fase 5/6: Divisione dell'audio in segmenti..."
        )
        chunks = await split_audio(wav_path)

        # Transcription
        update_job_status(
            job_id,
            "transcribing",
            current=current_step,
            total_chunks=total_steps,
            message=f"Fase 6/6: Trascrizione in corso (0/{len(chunks)} segmenti)"
        )

        # Batch processing ottimizzato per GPU
        for i, chunk in enumerate(chunks):
            try:
                # Forza sincronizzazione GPU prima di ogni chunk
                if torch.cuda.is_available():
                    torch.cuda.synchronize()
                    torch.cuda.empty_cache()

                with torch.cuda.amp.autocast(
                    enabled=True, dtype=torch.float16, cache_enabled=True
                ):
                    with torch.inference_mode():
                        with torch.backends.cuda.sdp_kernel(
                            enable_flash=True,
                            enable_math=True,
                            enable_mem_efficient=True,
                        ):
                            result = model.transcribe(
                                chunk,
                                fp16=True,
                                language="italian",
                                initial_prompt="Questa è una trascrizione di un audio in italiano.",
                                temperature=0.0,
                                no_speech_threshold=0.6,
                                compression_ratio_threshold=2.4,
                                condition_on_previous_text=True,
                                beam_size=5,  # Aumentato per sfruttare la GPU
                                best_of=5,  # Aumentato per sfruttare la GPU
                                without_timestamps=True,
                            )
                        transcriptions.append(result["text"])

                os.remove(chunk)

                # Forza sincronizzazione dopo ogni chunk
                if torch.cuda.is_available():
                    torch.cuda.synchronize()

                update_job_status(
                    job_id,
                    "transcribing",
                    i + 1,
                    len(chunks),
                    f"Trascrizione segmento {i+1} di {len(chunks)}"
                )

            except Exception as e:
                print(f"Errore nel chunk {i}: {str(e)}")

        # Fase 5: Formattazione
        current_step += 1
        update_job_status(
            job_id,
            "formatting",
            current=current_step,
            total_chunks=total_steps,
            message="Fase 5/5: Formattazione del documento finale..."
        )

        update_job_status(job_id, "formatting", message="Formattazione del testo...")
        full_text = "\n\n".join(transcriptions)
        formatted_text = format_transcription(full_text)

        update_job_status(
            job_id, "creating_document", message="Creazione del documento Word..."
        )

        # Crea il documento Word
        doc = Document()
        doc.add_heading("Trascrizione Audio", 0)
        doc.add_paragraph(
            f'Data trascrizione: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
        )
        doc.add_paragraph(f'File originale: {JOBS[job_id]["filename"]}')
        doc.add_paragraph("")

        for paragraph in formatted_text.split("\n"):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                p.paragraph_format.line_spacing = 1.5

        doc.save(output_path)

        JOBS[job_id]["status"] = "completed"
        JOBS[job_id]["output_path"] = output_path

    except Exception as e:
        print(f"Errore durante l'elaborazione del job {job_id}: {str(e)}")
        JOBS[job_id].update(
            {
                "status": "failed",
                "error": {
                    "message": str(e),
                    "type": type(e).__name__,
                    "details": "Si è verificato un errore durante l'elaborazione dell'audio",
                },
            }
        )
        raise
    finally:
        # Cleanup temporary files even if error occurs
        cleanup_temp_files(file_path)


def format_transcription(text: str) -> str:
    """Migliora la formattazione del testo trascritto"""
    import re

    # Correggi la punteggiatura comune
    text = re.sub(r"\s+([.,!?])", r"\1", text)
    text = re.sub(r"([.,!?])([A-Z])", r"\1\n\n\2", text)

    # Gestisci le pause nel parlato
    text = re.sub(r"\s*\.\.\.\s*", "...\n", text)

    # Aggiungi virgole prima delle congiunzioni
    conjunctions = r"\s+(e|ma|però|quindi|perché|quando|se)\s+"
    text = re.sub(conjunctions, lambda m: f", {m.group(1)} ", text)

    # Correggi la formattazione dei dialoghi
    text = re.sub(r'"([^"]+)"', r'"\1"', text)

    # Aggiungi spazi dopo la punteggiatura
    text = re.sub(r"([.,!?])([^\s])", r"\1 \2", text)

    return text


def cleanup_temp_files(file_path: str):
    """Clean up temporary files after processing"""
    try:
        # Clean up the original file
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Clean up WAV file if it exists
        wav_path = f"{os.path.splitext(file_path)[0]}.wav"
        if os.path.exists(wav_path):
            os.remove(wav_path)
            
        # Clean up any remaining chunk files
        base_path = os.path.splitext(file_path)[0]
        chunk_pattern = f"{base_path}_chunk_*.wav"
        for chunk_file in glob.glob(chunk_pattern):
            if os.path.exists(chunk_file):
                os.remove(chunk_file)
                
    except Exception as e:
        print(f"Error during cleanup: {str(e)}")


@app.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...), background_tasks: BackgroundTasks = BackgroundTasks()
):
    cleanup_old_jobs()

    # List of supported audio formats
    supported_formats = [".m4a", ".mp3", ".wav", ".aac", ".wma", ".ogg"]
    file_ext = os.path.splitext(file.filename)[1].lower()

    if file_ext not in supported_formats:
        raise HTTPException(
            status_code=400,
            detail=f"Formato file non supportato. Formati supportati: {', '.join(supported_formats)}",
        )

    # Generate unique job ID
    job_id = str(uuid.uuid4())

    try:
        # Create temp and output directories if they don't exist
        os.makedirs("temp", exist_ok=True)
        os.makedirs("output", exist_ok=True)

        # Save file
        file_path = f"temp/{job_id}_{file.filename}"
        output_path = f"output/{job_id}_{file.filename}.docx"

        with open(file_path, "wb") as buffer:
            # Read in chunks to handle large files
            chunk_size = 1024 * 1024  # 1MB chunks
            while chunk := await file.read(chunk_size):
                buffer.write(chunk)

        # Initialize job status
        JOBS[job_id] = {
            "status": "processing",
            "timestamp": datetime.now(),
            "filename": file.filename,
            "progress": {
                "stage": "initializing",
                "current_chunk": 0,
                "total_chunks": 0,
                "message": "Inizializzazione del processo...",
            },
        }

        # Start processing in background
        background_tasks.add_task(process_audio, job_id, file_path, output_path)

        return JSONResponse({"job_id": job_id, "status": "processing"})

    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/status/{job_id}")
async def get_job_status(job_id: str):
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="Job non trovato")
    return JOBS[job_id]


@app.get("/download/{job_id}")
async def download_file(job_id: str):
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="Job not found")

    job = JOBS[job_id]
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="File not ready")

    return FileResponse(job["output_path"], filename=f"{job['filename']}.docx")


async def load_model():
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
        torch.backends.cudnn.benchmark = True
        torch.backends.cuda.matmul.allow_tf32 = True
        torch.backends.cudnn.allow_tf32 = True

    model = whisper.load_model("small")

    if torch.cuda.is_available():
        model = model.cuda()
        model.eval()
        model = torch.compile(model, mode="max-autotune", fullgraph=True)

    return model
